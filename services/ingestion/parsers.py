"""
NEXUS — Document Parsers
Handles PDF, DOCX, TXT, Markdown parsing with OCR fallback.

Parsing is purely mechanical text extraction — no LLM involvement.
OCR is applied only when PDF pages yield sparse text.
"""
import io
from dataclasses import dataclass, field
from typing import Any


@dataclass
class ParsedDocument:
    """Result of parsing a document."""
    text: str
    page_count: int | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


def parse_document(
    *,
    file_data: bytes,
    mime_type: str,
    filename: str,
) -> ParsedDocument:
    """
    Parse a document based on its MIME type.
    Raises ValueError for unsupported types.
    """
    if mime_type == "application/pdf":
        return _parse_pdf(file_data, filename)
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _parse_docx(file_data, filename)
    elif mime_type in ("text/plain", "text/markdown"):
        return _parse_text(file_data, filename)
    else:
        raise ValueError(f"Unsupported MIME type: {mime_type}")


def _parse_pdf(file_data: bytes, filename: str) -> ParsedDocument:
    """Parse PDF using pdfplumber. Falls back to OCR if text is sparse."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber is required for PDF parsing.")

    pages_text: list[str] = []
    metadata: dict = {}

    with pdfplumber.open(io.BytesIO(file_data)) as pdf:
        metadata = {
            "title": pdf.metadata.get("Title"),
            "author": pdf.metadata.get("Author"),
            "creator": pdf.metadata.get("Creator"),
            "creation_date": str(pdf.metadata.get("CreationDate", "")),
        }
        page_count = len(pdf.pages)

        for page in pdf.pages:
            text = page.extract_text() or ""
            # If a page has very little text, attempt OCR
            if len(text.strip()) < 50:
                ocr_text = _ocr_pdf_page(page)
                if ocr_text:
                    text = ocr_text
            pages_text.append(text)

    full_text = "\n\n".join(pages_text)
    metadata["page_count"] = page_count
    metadata["parser"] = "pdfplumber"

    return ParsedDocument(
        text=_clean_text(full_text),
        page_count=page_count,
        metadata={k: v for k, v in metadata.items() if v},
    )


def _ocr_pdf_page(page) -> str:
    """OCR a single PDF page using pytesseract."""
    try:
        import pytesseract
        img = page.to_image(resolution=200).original
        return pytesseract.image_to_string(img, config="--psm 6")
    except Exception:
        return ""


def _parse_docx(file_data: bytes, filename: str) -> ParsedDocument:
    """Parse DOCX using python-docx."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx is required for DOCX parsing.")

    doc = DocxDocument(io.BytesIO(file_data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)

    metadata: dict = {"parser": "python-docx"}
    try:
        core = doc.core_properties
        if core.title:
            metadata["title"] = core.title
        if core.author:
            metadata["author"] = core.author
    except Exception:
        pass

    return ParsedDocument(
        text=_clean_text(full_text),
        page_count=None,
        metadata=metadata,
    )


def _parse_text(file_data: bytes, filename: str) -> ParsedDocument:
    """Parse plain text or Markdown."""
    text = file_data.decode("utf-8", errors="replace")
    return ParsedDocument(
        text=_clean_text(text),
        page_count=None,
        metadata={"parser": "text", "encoding": "utf-8"},
    )


def _clean_text(text: str) -> str:
    """
    Normalize whitespace and remove control characters.
    Does NOT alter content — no summarization or modification.
    """
    import re
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
